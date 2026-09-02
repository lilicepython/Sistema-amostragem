import streamlit as st
import sqlite3
import io
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# =====================================================================
# INICIALIZAÇÃO DO BANCO DE DADOS
# =====================================================================
def iniciar_banco():
    conexao = sqlite3.connect('laboratorio.db')
    cursor = conexao.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS ordens_servico_arquivos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT,
            tamanho TEXT,
            data_anexo TEXT,
            conteudo_pdf BLOB
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS cadeias_custodia_geradas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            os_num TEXT,
            empreendimento TEXT,
            matriz TEXT,
            data_geracao TEXT,
            conteudo_pdf BLOB,
            conteudo_docx BLOB
        )
    ''')
    conexao.commit()
    conexao.close()

iniciar_banco()

def deletar_os(id_registro):
    conexao = sqlite3.connect('laboratorio.db')
    cursor = conexao.cursor()
    cursor.execute("DELETE FROM ordens_servico_arquivos WHERE id = ?", (id_registro,))
    conexao.commit()
    conexao.close()

def deletar_cc(id_registro):
    conexao = sqlite3.connect('laboratorio.db')
    cursor = conexao.cursor()
    cursor.execute("DELETE FROM cadeias_custodia_geradas WHERE id = ?", (id_registro,))
    conexao.commit()
    conexao.close()

# =====================================================================
# CONFIGURAÇÃO GERAL
# =====================================================================
st.set_page_config(page_title="Sistema de Amostragem", layout="wide")
st.markdown("""
    <style>
    :root { --azul-neon: #00FFFF; --azul-neon-escuro: #00cccc; }
    .stButton>button { background-color: var(--azul-neon); color: black; font-weight: bold; border-radius: 5px; }
    .stButton>button:hover { background-color: var(--azul-neon-escuro); color: white; }
    h1, h2, h3 { color: var(--azul-neon); }
    </style>
""", unsafe_allow_html=True)

# =====================================================================
# NAVEGAÇÃO LATERAL
# =====================================================================
st.sidebar.title("Navegação do Sistema")
aba_selecionada = st.sidebar.radio(
    "Selecione o Módulo:",
    ["Ordem de Serviço", "Plano de Amostragem", "Cadeia de Custódia"]
)
st.sidebar.markdown("---")

# =====================================================================
# MÓDULO 1: ORDEM DE SERVIÇO
# =====================================================================
if aba_selecionada == "Ordem de Serviço":
    st.title("Gestão de Ordens de Serviço")
    st.write("Anexe propostas comerciais em formato PDF.")
    
    arquivo_upload = st.file_uploader("Carregar Proposta/OS (Somente PDF)", type=["pdf"])
    
    if st.button("Salvar Arquivo"):
        if arquivo_upload:
            nome = arquivo_upload.name
            tamanho = f"{arquivo_upload.size / 1024:.2f} KB"
            data_anexo = date.today().strftime("%d/%m/%Y")
            conteudo = arquivo_upload.getvalue()
            
            conexao = sqlite3.connect('laboratorio.db')
            cursor = conexao.cursor()
            cursor.execute('INSERT INTO ordens_servico_arquivos (nome, tamanho, data_anexo, conteudo_pdf) VALUES (?, ?, ?, ?)', 
                           (nome, tamanho, data_anexo, conteudo))
            conexao.commit()
            conexao.close()
            st.success(f"Arquivo '{nome}' salvo com sucesso!")
            st.rerun()
        else:
            st.warning("Selecione um arquivo PDF.")
            
    st.markdown("---")
    st.subheader("Histórico de Arquivos Anexados")
    conexao = sqlite3.connect('laboratorio.db')
    cursor = conexao.cursor()
    cursor.execute("SELECT id, nome, tamanho, data_anexo, conteudo_pdf FROM ordens_servico_arquivos ORDER BY id DESC")
    arquivos = cursor.fetchall()
    conexao.close()
    
    if arquivos:
        for arq in arquivos:
            id_arq, nome_arq, tam_arq, data_arq, pdf_bytes = arq
            c1, c2, c3 = st.columns([6, 2, 2])
            with c1:
                st.markdown(f"**{nome_arq}** | {tam_arq} | {data_arq}")
            with c2:
                st.download_button("Baixar PDF", data=pdf_bytes, file_name=nome_arq, mime="application/pdf", key=f"dl_os_{id_arq}")
            with c3:
                st.button("Excluir", key=f"del_os_{id_arq}", on_click=deletar_os, args=(id_arq,))
            st.markdown("---")

# =====================================================================
# MÓDULO 2: PLANO DE AMOSTRAGEM
# =====================================================================
elif aba_selecionada == "Plano de Amostragem":
    st.title("Geração de Planos de Amostragem")
    
    num_documento = st.text_input("Numeração do Documento")
    nome_empreendimento = st.text_input("Nome do Empreendimento")
    endereco_empreendimento = st.text_input("Endereço do Empreendimento")
    responsavel_empreendimento = st.text_input("Responsável")
    
    if st.button("Gerar Plano de Amostragem", type="primary"):
        if not num_documento or not nome_empreendimento:
            st.error("Preencha ao menos a Numeração e o Empreendimento.")
        else:
            # Word
            doc = Document()
            titulo = doc.add_heading('PLANO DE AMOSTRAGEM', 0)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            t = doc.add_table(rows=4, cols=2)
            t.style = 'Table Grid'
            t.cell(0,0).text, t.cell(0,1).text = "Documento N°:", num_documento
            t.cell(1,0).text, t.cell(1,1).text = "Empreendimento:", nome_empreendimento
            t.cell(2,0).text, t.cell(2,1).text = "Endereço:", endereco_empreendimento
            t.cell(3,0).text, t.cell(3,1).text = "Responsável:", responsavel_empreendimento
            buf_docx = io.BytesIO()
            doc.save(buf_docx)
            
            # PDF
            buf_pdf = io.BytesIO()
            pdf = SimpleDocTemplate(buf_pdf, pagesize=A4)
            estilos = getSampleStyleSheet()
            elementos = [Paragraph("PLANO DE AMOSTRAGEM", ParagraphStyle(name='Title', parent=estilos['Heading1'], alignment=1)), Spacer(1, 15)]
            dados_tabela = [
                ["Documento N°:", num_documento], ["Empreendimento:", nome_empreendimento],
                ["Endereço:", endereco_empreendimento], ["Responsável:", responsavel_empreendimento]
            ]
            tabela = Table(dados_tabela, colWidths=[150, 350])
            tabela.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.black), ('BACKGROUND', (0,0), (0,-1), colors.lightgrey)]))
            elementos.append(tabela)
            pdf.build(elementos)
            
            c1, c2 = st.columns(2)
            c1.download_button("Baixar Plano (PDF)", data=buf_pdf.getvalue(), file_name=f"Plano_{num_documento}.pdf", mime="application/pdf")
            c2.download_button("Baixar Plano (Word)", data=buf_docx.getvalue(), file_name=f"Plano_{num_documento}.docx")

# =====================================================================
# MÓDULO 3: CADEIA DE CUSTÓDIA
# =====================================================================
elif aba_selecionada == "Cadeia de Custódia":
    st.title("Gestão de Cadeias de Custódia")
    tab_gerar, tab_hist = st.tabs(["Gerar Cadeia de Custódia", "Cadeias de Custódia Geradas"])
    
    with tab_gerar:
        st.subheader("Identificação do Empreendimento")
        c1, c2, c3 = st.columns(3)
        empreendimento = c1.text_input("Empreendimento")
        endereco = c1.text_input("Endereço")
        cod_cliente = c2.text_input("Cód. Cliente")
        responsavel = c2.text_input("Responsável")
        os_num = c3.text_input("OS N° / PC N°")
        contato = c3.text_input("Contato")
        ref = st.text_input("Ponto de Referência")

        st.subheader("Identificação da Amostragem")
        c_m1, c_m2 = st.columns(2)
        opcoes_matriz = {
            "Água e Efluentes": ["ACH", "ASP", "ASB", "EFL"],
            "Emissões Atmosféricas (EAT)": ["Isocinética", "Gases de Combustão", "Fumaça"],
            "Ruído e Vibração": ["Acústica (ACT)", "Dosimetria (DOS)"],
            "Qualidade do Ar": ["ATM (Gases)", "QAI (Bactérias/Fungos)"],
            "Resíduos Sólidos (RS)": ["Sólido", "Líquido"],
            "Solo": ["Solo Padrão"],
            "Sedimentos": ["Sedimento Padrão"]
        }
        matriz = c_m1.selectbox("Matriz", list(opcoes_matriz.keys()))
        submatriz = c_m2.selectbox("Submatriz", opcoes_matriz[matriz])
        
        c_d1, c_d2 = st.columns(2)
        data_coleta = c_d1.date_input("Data da Coleta", format="DD/MM/YYYY")
        hora_coleta = c_d2.time_input("Hora da Coleta (Opcional)")
        
        hoje = date.today()
        bloqueio_temporal = hoje < data_coleta
        
        if bloqueio_temporal:
            st.error(f"⚠️ Acesso Bloqueado: A Data da Coleta ({data_coleta.strftime('%d/%m/%Y')}) é no futuro. O preenchimento in loco e a emissão documental estão suspensos.")
        else:
            st.success("✅ Acesso Liberado: Data validada.")
            
        st.markdown("---")
        st.subheader("Parâmetros das medições in loco")
        qtd_pontos = st.number_input("Quantidade de Pontos (Máx: 100)", min_value=1, max_value=100, step=1, disabled=bloqueio_temporal)
        
        dados_pontos = []
        for i in range(int(qtd_pontos)):
            with st.expander(f"Ponto {i+1}", expanded=not bloqueio_temporal):
                p1, p2, p3 = st.columns(3)
                id_ponto = p1.text_input("ID do Ponto", key=f"id_{i}", disabled=bloqueio_temporal)
                ph = p1.text_input("pH", key=f"ph_{i}", disabled=bloqueio_temporal)
                temp = p1.text_input("Temp (°C)", key=f"t_{i}", disabled=bloqueio_temporal)
                cond = p2.text_input("Condutividade", key=f"c_{i}", disabled=bloqueio_temporal)
                od = p2.text_input("Oxigênio Dissolvido", key=f"o_{i}", disabled=bloqueio_temporal)
                std = p2.text_input("STD", key=f"st_{i}", disabled=bloqueio_temporal)
                sal = p3.text_input("Salinidade", key=f"sa_{i}", disabled=bloqueio_temporal)
                res = p3.text_input("Resistividade", key=f"re_{i}", disabled=bloqueio_temporal)
                orp = p3.text_input("ORP", key=f"or_{i}", disabled=bloqueio_temporal)
                dados_pontos.append([id_ponto, ph, temp, cond, od, std, sal, res, orp])

        st.markdown("---")
        st.subheader("Recepção e Inspeção da Amostra")
        r1, r2, r3 = st.columns(3)
        entregue = r1.text_input("Entregue por", disabled=bloqueio_temporal)
        recebido = r1.text_input("Recebido por", disabled=bloqueio_temporal)
        data_rec = r2.date_input("Data Recepção", format="DD/MM/YYYY", disabled=bloqueio_temporal)
        hora_rec = r2.time_input("Hora Recepção", disabled=bloqueio_temporal)
        temp_rec = r3.text_input("Temperatura Recepção", disabled=bloqueio_temporal)
        desvio = r3.text_input("Desvio? (Sim/Não)", disabled=bloqueio_temporal)
        
        if not bloqueio_temporal:
            if st.button("Gerar Cadeia de Custódia", type="primary"):
                orientacao = landscape(A4) if matriz == "Emissões Atmosféricas (EAT)" else A4
                
                # DOCX
                doc = Document()
                if orientacao == landscape(A4):
                    secao = doc.sections[-1]
                    secao.orientation = WD_ORIENT.LANDSCAPE
                    secao.page_width, secao.page_height = secao.page_height, secao.page_width
                    
                tit = doc.add_heading('CADEIA DE CUSTÓDIA', 0)
                tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_heading('Identificação do Empreendimento', 1)
                t_emp = doc.add_table(rows=3, cols=2)
                t_emp.style = 'Table Grid'
                t_emp.cell(0,0).text = f"Empreendimento: {empreendimento}"
                t_emp.cell(0,1).text = f"Cód/OS: {cod_cliente} / {os_num}"
                t_emp.cell(1,0).text = f"Endereço: {endereco}"
                t_emp.cell(1,1).text = f"Responsável: {responsavel}"
                t_emp.cell(2,0).text = f"Referência: {ref}"
                t_emp.cell(2,1).text = f"Contato: {contato}"
                
                doc.add_heading('Identificação da Amostragem', 1)
                t_amo = doc.add_table(rows=1, cols=3)
                t_amo.style = 'Table Grid'
                t_amo.cell(0,0).text = f"Matriz: {matriz}"
                t_amo.cell(0,1).text = f"Submatriz: {submatriz}"
                t_amo.cell(0,2).text = f"Data: {data_coleta.strftime('%d/%m/%Y')} {hora_coleta.strftime('%H:%M')}"
                
                doc.add_heading('Parâmetros das medições in loco', 1)
                cabs = ["Ponto", "pH", "Temp", "Cond", "OD", "STD", "Sal", "Resist", "ORP"]
                t_par = doc.add_table(rows=1, cols=len(cabs))
                t_par.style = 'Table Grid'
                for i, c in enumerate(cabs): t_par.cell(0,i).text = c
                
                # Tabela gerada mesmo com campos em branco garantindo o grid para preenchimento posterior
                for linha in dados_pontos:
                    rc = t_par.add_row().cells
                    for i, val in enumerate(linha): rc[i].text = val
                
                doc.add_heading('Recepção e Inspeção', 1)
                t_rec = doc.add_table(rows=2, cols=2)
                t_rec.style = 'Table Grid'
                t_rec.cell(0,0).text = f"Entregue por: {entregue}"
                t_rec.cell(0,1).text = f"Recebido por: {recebido}"
                t_rec.cell(1,0).text = f"Data: {data_rec.strftime('%d/%m/%Y')} {hora_rec.strftime('%H:%M')}"
                t_rec.cell(1,1).text = f"Temp: {temp_rec} | Desvio: {desvio}"
                
                buf_docx = io.BytesIO()
                doc.save(buf_docx)

                # PDF
                buf_pdf = io.BytesIO()
                pdf = SimpleDocTemplate(buf_pdf, pagesize=orientacao)
                estilos = getSampleStyleSheet()
                est_titulo = ParagraphStyle('tit', parent=estilos['Heading1'], alignment=1)
                
                elem = [Paragraph("CADEIA DE CUSTÓDIA", est_titulo), Spacer(1, 10)]
                est_tab = TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.black), ('BACKGROUND', (0,0), (-1,0), colors.lightgrey)])
                
                elem.append(Paragraph("Identificação", estilos['Heading3']))
                elem.append(Table([[f"Emp: {empreendimento}", f"OS: {os_num}"], [f"End: {endereco}", f"Resp: {responsavel}"]], style=est_tab))
                
                elem.append(Paragraph("Amostragem", estilos['Heading3']))
                elem.append(Table([[matriz, submatriz, f"Data: {data_coleta.strftime('%d/%m/%Y')}"]], style=est_tab))
                
                elem.append(Paragraph("Parâmetros in loco", estilos['Heading3']))
                tabela_pontos = Table([cabs] + dados_pontos, repeatRows=1)
                tabela_pontos.setStyle(est_tab)
                elem.append(tabela_pontos)
                
                elem.append(Paragraph("Recepção", estilos['Heading3']))
                elem.append(Table([[f"Entregue: {entregue}", f"Recebido: {recebido}"], [f"Data: {data_rec.strftime('%d/%m/%Y')}", f"Temp: {temp_rec}"]], style=est_tab))
                
                pdf.build(elem)
                
                # Salvar DB
                conexao = sqlite3.connect('laboratorio.db')
                c = conexao.cursor()
                c.execute('INSERT INTO cadeias_custodia_geradas (os_num, empreendimento, matriz, data_geracao, conteudo_pdf, conteudo_docx) VALUES (?,?,?,?,?,?)',
                          (os_num, empreendimento, matriz, date.today().strftime("%d/%m/%Y"), buf_pdf.getvalue(), buf_docx.getvalue()))
                conexao.commit()
                conexao.close()
                st.success("Cadeia de Custódia gerada com sucesso! Acesse a aba 'Cadeias de Custódia Geradas'.")

    with tab_hist:
        st.subheader("Cadeias de Custódia Salvas")
        conexao = sqlite3.connect('laboratorio.db')
        c = conexao.cursor()
        c.execute("SELECT id, os_num, empreendimento, matriz, data_geracao, conteudo_pdf, conteudo_docx FROM cadeias_custodia_geradas")
        salvos = c.fetchall()
        conexao.close()
        
        if salvos:
            for cc in salvos:
                id_cc, os_db, emp_db, mat_db, data_db, pdf_b, doc_b = cc
                c1, c2, c3 = st.columns([5, 2, 2])
                with c1: st.write(f"**OS {os_db}** | {emp_db} ({mat_db}) | {data_db}")
                with c2: st.download_button("PDF", pdf_b, f"CC_{os_db}.pdf", key=f"pdf_{id_cc}")
                with c3: st.download_button("Word", doc_b, f"CC_{os_db}.docx", key=f"doc_{id_cc}")
                st.button("Excluir", key=f"del_cc_{id_cc}", on_click=deletar_cc, args=(id_cc,))
                st.markdown("---")